# streamlit_app.py

import streamlit as st
from io import BytesIO
from docx import Document

from run_pipeline import execute_query


# --------------------------------------------------
# Helper: Word Export
# --------------------------------------------------

def create_word_file(content):
    """
    Create downloadable Word (.docx) file
    """

    doc = Document()
    doc.add_heading(
        "Executive Summary",
        level=1
    )

    doc.add_paragraph(content)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


# --------------------------------------------------
# Page Config
# --------------------------------------------------

st.set_page_config(
    page_title="AI Supply Chain Strategic Copilot",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)


# --------------------------------------------------
# Optional Styling
# --------------------------------------------------

st.markdown(
    """
    <style>
        .main {
            padding-top: 1rem;
        }

        .stButton > button {
            width: 100%;
            height: 3rem;
            font-size: 16px;
            font-weight: 600;
            border-radius: 10px;
        }

        .block-container {
            padding-top: 2rem;
            padding-bottom: 2rem;
        }
    </style>
    """,
    unsafe_allow_html=True
)


# --------------------------------------------------
# Sidebar
# --------------------------------------------------

with st.sidebar:
    st.image(
        "logo.png",
        width=220
    )

    st.title(
        "Strategic Copilot"
    )

    st.markdown("---")

    st.write("### Platform Overview")

    st.write(
        """
        Enterprise AI system for:

        - Inventory risk analysis
        - Product ranking
        - Supply chain diagnostics
        - Business comparison
        - Executive decision support
        """
    )

    st.markdown("---")

    st.write("### Core Engine")

    st.write(
        """
        - LangGraph Workflow
        - Query Validator
        - Self-Healing Agent
        - DSL Executor
        - Reflection Layer
        """
    )


# --------------------------------------------------
# Main Header
# --------------------------------------------------

st.title(
    "AI Supply Chain Strategic Copilot"
)

st.caption(
    "Executive-grade analytics for supply chain intelligence"
)

st.markdown("---")


# --------------------------------------------------
# Example Queries
# --------------------------------------------------

with st.expander(
    "Example Strategic Questions"
):
    st.write("- Top 10 risky products")
    st.write("- Why is TX_3 risky?")
    st.write("- Compare CA_1 vs TX_2")
    st.write("- Which category has highest inventory risk?")
    st.write("- Which state has margin erosion?")


# --------------------------------------------------
# User Input
# --------------------------------------------------

user_input = st.text_input(
    "Ask your strategic question:",
    placeholder="Example: Top 10 risky products"
)


# --------------------------------------------------
# Run Button
# --------------------------------------------------

if st.button(
    "Run Strategic Analysis"
):

    if not user_input.strip():
        st.warning(
            "Please enter a valid business question."
        )
        st.stop()

    with st.spinner(
        "Running full strategic pipeline..."
    ):
        try:
            final_state = execute_query(
                user_input
            )

            st.markdown("---")

            # -----------------------------------
            # Failure Path
            # -----------------------------------

            if final_state.get(
                "failure_stage"
            ):

                st.error(
                    "System entered reflection mode"
                )

                st.subheader(
                    "Reflection Output"
                )

                st.write(
                    final_state.get(
                        "reflection_output",
                        "No reflection available."
                    )
                )

            # -----------------------------------
            # Success Path
            # -----------------------------------

            else:
                executive_summary = final_state.get(
                    "final_response",
                    "No final response generated."
                )

                st.success(
                    "Analysis completed successfully"
                )

                st.subheader(
                    "Executive Summary"
                )

                st.write(
                    executive_summary
                )

                # ------------------------------
                # Confidence Score
                # ------------------------------

                confidence = final_state.get(
                    "confidence_score",
                    0.0
                )

                st.metric(
                    label="Confidence Score",
                    value=f"{round(confidence * 100, 1)}%"
                )

                # ------------------------------
                # Export Word File
                # ------------------------------

                word_file = create_word_file(
                    executive_summary
                )

                st.download_button(
                    label="Export Executive Summary (.docx)",
                    data=word_file,
                    file_name="executive_summary.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            # -----------------------------------
            # Technical Debug View
            # -----------------------------------

            with st.expander(
                "Technical Debug View"
            ):
                st.json(
                    final_state
                )

        except Exception as e:
            st.error(
                f"System Error: {str(e)}"
            )